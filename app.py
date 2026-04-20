import streamlit as st
import os
from datetime import date
from supabase import create_client
from dotenv import load_dotenv
from job_agent import build_graph, AgentState
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO

load_dotenv()

# ---- SUPABASE SETUP ----
supabase = create_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_KEY")
)

# ---- PAGE CONFIG ----
st.set_page_config(
    page_title="ProofHire — GitHub-Verified Job Matching",
    page_icon="🎯",
    layout="wide"
)

# ---- HIDE SIDEBAR + ALL CSS ----
st.markdown("""
    <style>
    [data-testid="stSidebar"] { display: none !important; }
    [data-testid="stSidebarNav"] { display: none !important; }
    #MainMenu { visibility: hidden; }
    header { visibility: hidden; }
    footer { visibility: hidden; }
    .block-container {
        padding-top: 80px !important;
        padding-left: 40px !important;
        padding-right: 40px !important;
        max-width: 1400px !important;
    }
    .main { background-color: #f8f9fa; }

    /* NAVBAR */
    .navbar {
        position: fixed;
        top: 0; left: 0; right: 0;
        height: 64px;
        background: white;
        border-bottom: 1px solid #e2e8f0;
        display: flex;
        align-items: center;
        justify-content: space-between;
        padding: 0 40px;
        z-index: 999;
        box-shadow: 0 1px 8px rgba(0,0,0,0.05);
    }
    .navbar-logo {
        display: flex;
        align-items: center;
        gap: 12px;
    }
    .logo-box {
        width: 40px; height: 40px;
        background: #2563eb;
        border-radius: 10px;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    .logo-text {
        font-size: 1.4rem;
        font-weight: 900;
        color: #0f172a;
        letter-spacing: -0.5px;
        line-height: 1;
    }
    .logo-text span { color: #2563eb; }
    .logo-tagline {
        font-size: 0.7rem;
        color: #94a3b8;
        margin-top: 2px;
        letter-spacing: 0.3px;
    }
    .navbar-right {
        display: flex;
        align-items: center;
        gap: 4px;
    }
    .nav-divider {
        width: 1px; height: 22px;
        background: #e2e8f0;
        margin: 0 10px;
    }
    .nav-by {
        font-size: 0.75rem;
        color: #94a3b8;
    }
    .nav-by strong { color: #2563eb; }

    /* HERO */
    .hero-container {
        background: linear-gradient(135deg, #ffffff 0%, #f0f4ff 100%);
        border: 1px solid #e0e7ff;
        border-radius: 16px;
        padding: 50px 40px;
        text-align: center;
        margin-bottom: 30px;
    }
    .hero-badge {
        background-color: #eff6ff;
        color: #2563eb;
        border: 1px solid #bfdbfe;
        border-radius: 20px;
        padding: 4px 14px;
        font-size: 0.8rem;
        font-weight: 600;
        display: inline-block;
        margin-bottom: 16px;
    }
    .hero-title {
        font-size: 3rem;
        font-weight: 800;
        color: #0f172a;
        margin-bottom: 12px;
        letter-spacing: -1px;
    }
    .hero-title span { color: #2563eb; }
    .hero-subtitle {
        font-size: 1.1rem;
        color: #64748b;
        max-width: 600px;
        margin: 0 auto 24px auto;
        line-height: 1.7;
    }
    .hero-stats {
        display: flex;
        justify-content: center;
        gap: 40px;
        margin-top: 30px;
    }
    .stat-item { text-align: center; }
    .stat-number { font-size: 1.8rem; font-weight: 800; color: #2563eb; }
    .stat-label {
        font-size: 0.8rem;
        color: #94a3b8;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }

    /* AGENT CARDS */
    .agent-card {
        background: white;
        border: 1px solid #e2e8f0;
        border-radius: 12px;
        padding: 24px;
        text-align: center;
        height: 100%;
    }
    .agent-card:hover {
        border-color: #2563eb;
        box-shadow: 0 4px 20px rgba(37,99,235,0.1);
    }
    .agent-icon { font-size: 2rem; margin-bottom: 12px; }
    .agent-number {
        font-size: 0.75rem;
        color: #2563eb;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 6px;
    }
    .agent-title { font-size: 1rem; font-weight: 700; color: #0f172a; margin-bottom: 8px; }
    .agent-desc { font-size: 0.85rem; color: #64748b; line-height: 1.5; }

    /* PILLS */
    .feature-pill {
        background: #eff6ff;
        color: #2563eb;
        border-radius: 20px;
        padding: 6px 16px;
        font-size: 0.85rem;
        font-weight: 500;
        display: inline-block;
        margin: 4px;
    }

    /* SECTION */
    .section-header { font-size: 1.4rem; font-weight: 700; color: #0f172a; margin-bottom: 4px; }
    .section-sub { font-size: 0.9rem; color: #94a3b8; margin-bottom: 20px; }

    /* BUTTONS */
    .stButton > button {
        background: linear-gradient(135deg, #2563eb, #1d4ed8);
        color: white;
        border: none;
        border-radius: 8px;
        font-weight: 600;
        font-size: 1rem;
        padding: 12px 24px;
        width: 100%;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, #1d4ed8, #1e40af);
        box-shadow: 0 4px 15px rgba(37,99,235,0.3);
    }

    /* RESULTS */
    .result-header {
        background: linear-gradient(135deg, #2563eb, #7c3aed);
        color: white;
        border-radius: 12px;
        padding: 20px 24px;
        margin-bottom: 20px;
    }
    .result-title { font-size: 1.3rem; font-weight: 700; margin-bottom: 4px; }
    .result-sub { font-size: 0.85rem; opacity: 0.8; }

    /* FOOTER */
    .footer {
        text-align: center;
        color: #94a3b8;
        font-size: 0.8rem;
        padding: 20px 0;
        border-top: 1px solid #e2e8f0;
        margin-top: 40px;
    }

    /* POWERED BY PILLS */
    .tech-pill {
        display: inline-flex;
        align-items: center;
        gap: 6px;
        padding: 6px 12px;
        border-radius: 8px;
        font-size: 0.8rem;
        font-weight: 500;
        margin: 3px;
    }
    </style>
""", unsafe_allow_html=True)

# ---- WORD DOC HELPER ----
def create_word_doc(cover_letter, user_name, company_name):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    for line in cover_letter.split('\n'):
        line = line.strip()
        if line == '---':
            continue
        elif line == '':
            doc.add_paragraph('')
        else:
            para = doc.add_paragraph(line)
            run = para.runs[0] if para.runs else para.add_run(line)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---- TOP NAVBAR ----
st.markdown("""
    <div class="navbar">
        <div class="navbar-logo">
            <div class="logo-box">
                <svg width="22" height="22" viewBox="0 0 22 22" fill="none">
                    <circle cx="11" cy="11" r="9" stroke="white" stroke-width="1.8"/>
                    <path d="M7 11 L10 14 L15 8" stroke="white" stroke-width="2"
                          stroke-linecap="round" stroke-linejoin="round"/>
                </svg>
            </div>
            <div>
                <div class="logo-text">Proof<span>Hire</span></div>
                <div class="logo-tagline">GitHub-Verified Job Matching</div>
            </div>
        </div>
        <div class="navbar-right" id="nav-links">
        </div>
    </div>
""", unsafe_allow_html=True)

# ---- NAVIGATION ----
col_nav1, col_nav2, col_nav3, col_spacer, col_by = st.columns([1, 1.5, 1, 4, 1.5])

with col_nav1:
    home_btn = st.button("🏠 Home", use_container_width=True)
with col_nav2:
    gen_btn = st.button("📝 Generate Application", use_container_width=True)
with col_nav3:
    hist_btn = st.button("📚 My History", use_container_width=True)
with col_by:
    st.markdown("""
        <div style="font-size:0.75rem; color:#94a3b8; padding-top:8px; text-align:right;">
            by <strong style="color:#2563eb;">Harsha Bodhe</strong>
        </div>
    """, unsafe_allow_html=True)

# ---- PAGE STATE ----
if "page" not in st.session_state:
    st.session_state.page = "🏠 Home"

if home_btn:
    st.session_state.page = "🏠 Home"
if gen_btn:
    st.session_state.page = "📝 Generate Application"
if hist_btn:
    st.session_state.page = "📚 My History"

page = st.session_state.page

st.markdown("---")

# ============================
# PAGE 1 — HOME
# ============================
if page == "🏠 Home":

    st.markdown("""
        <div class="hero-container">
            <div class="hero-badge">🚀 4-Agent AI System · Powered by LangGraph</div>
            <div class="hero-title">Stop Sending Generic CVs.<br>Start <span>Proving</span> Your Skills.</div>
            <div class="hero-subtitle">
                ProofHire analyses your actual GitHub projects, matches them against job requirements,
                and generates evidence-based cover letters that get noticed.
            </div>
            <div>
                <span class="feature-pill">🐙 GitHub Verified</span>
                <span class="feature-pill">🎯 Fit Score</span>
                <span class="feature-pill">✉️ Cover Letter</span>
                <span class="feature-pill">📊 Skill Gap Analysis</span>
                <span class="feature-pill">📁 Word Export</span>
            </div>
            <div class="hero-stats">
                <div class="stat-item">
                    <div class="stat-number">4</div>
                    <div class="stat-label">AI Agents</div>
                </div>
                <div class="stat-item">
                    <div class="stat-number">100%</div>
                    <div class="stat-label">Evidence Based</div>
                </div>
                <div class="stat-item">
                    <div class="stat-number">&lt;60s</div>
                    <div class="stat-label">Per Application</div>
                </div>
            </div>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("### How It Works")
    st.markdown(
        "<p style='color:#64748b; margin-bottom:20px;'>"
        "Four specialised AI agents collaborate to build your perfect application</p>",
        unsafe_allow_html=True
    )

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown("""
            <div class="agent-card">
                <div class="agent-icon">🔍</div>
                <div class="agent-number">Agent 01</div>
                <div class="agent-title">GitHub Analyser</div>
                <div class="agent-desc">Fetches your real repositories and verifies
                your actual technical skills from code — not just your CV claims</div>
            </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown("""
            <div class="agent-card">
                <div class="agent-icon">📄</div>
                <div class="agent-number">Agent 02</div>
                <div class="agent-title">CV Analyser</div>
                <div class="agent-desc">Extracts your experience, education, projects
                and achievements from your uploaded CV or pasted text</div>
            </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown("""
            <div class="agent-card">
                <div class="agent-icon">🎯</div>
                <div class="agent-number">Agent 03</div>
                <div class="agent-title">Job Matcher</div>
                <div class="agent-desc">Calculates your Technical Fit Score, identifies
                your top matching skills and highlights skill gaps to address</div>
            </div>
        """, unsafe_allow_html=True)
    with col4:
        st.markdown("""
            <div class="agent-card">
                <div class="agent-icon">✍️</div>
                <div class="agent-number">Agent 04</div>
                <div class="agent-title">Cover Letter Writer</div>
                <div class="agent-desc">Generates a tailored, German-standard cover letter
                using GitHub evidence and your match analysis</div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Tech stack pills
    st.markdown("### Powered By")
    st.markdown("""
        <div style="display:flex; flex-wrap:wrap; gap:8px; margin-bottom:24px;">
            <span class="tech-pill" style="background:#f0f4ff; color:#2563eb; border:1px solid #bfdbfe;">
                🤖 LangGraph Multi-Agent
            </span>
            <span class="tech-pill" style="background:#f0fdf4; color:#16a34a; border:1px solid #bbf7d0;">
                ⚡ Groq LLM Inference
            </span>
            <span class="tech-pill" style="background:#fef3c7; color:#d97706; border:1px solid #fcd34d;">
                🐙 GitHub API Verification
            </span>
            <span class="tech-pill" style="background:#fdf4ff; color:#7c3aed; border:1px solid #e9d5ff;">
                🗄️ Supabase Storage
            </span>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="background: linear-gradient(135deg, #eff6ff, #f0fdf4);
                    border: 1px solid #bfdbfe; border-radius: 12px;
                    padding: 24px; text-align: center;">
            <div style="font-size:1.1rem; font-weight:700; color:#0f172a; margin-bottom:8px;">
                Ready to get started?
            </div>
            <div style="color:#64748b; font-size:0.9rem;">
                Click <strong>Generate Application</strong> in the navbar above
            </div>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div class="footer">
            🎯 ProofHire · GitHub-Verified Job Matching ·
            Powered by LangGraph + Groq + Supabase · 2026
        </div>
    """, unsafe_allow_html=True)

# ============================
# PAGE 2 — GENERATE APPLICATION
# ============================
elif page == "📝 Generate Application":

    st.markdown("""
        <div style="margin-bottom:24px;">
            <div class="section-header">📝 Generate Your Application</div>
            <div class="section-sub">Fill in your details and let ProofHire's 4 agents do the work</div>
        </div>
    """, unsafe_allow_html=True)

    st.subheader("👤 Your Details")
    col1, col2 = st.columns(2)
    with col1:
        user_name = st.text_input("Full Name *", placeholder="Max Müller")
        user_email = st.text_input("Email *", placeholder="max.muller@gmail.com")
        user_phone = st.text_input("Phone Number", placeholder="+49 155 12345678")
    with col2:
        user_location = st.text_input("City, Country", placeholder="Berlin, Germany")
        user_linkedin = st.text_input("LinkedIn URL", placeholder="linkedin.com/in/maxmuller")
        github_username = st.text_input("GitHub Username * (no spaces)", placeholder="MaxMuller")

    st.markdown("---")

    st.markdown("**📄 Your CV**")
    cv_upload_method = st.radio(
        "How would you like to provide your CV?",
        ["📁 Upload PDF or Word file", "✏️ Paste text manually"],
        horizontal=True
    )

    cv_text = ""

    if cv_upload_method == "📁 Upload PDF or Word file":
        uploaded_file = st.file_uploader(
            "Upload your CV",
            type=["pdf", "docx", "txt"],
            help="Supports PDF, Word (.docx) and text files"
        )
        if uploaded_file is not None:
            file_type = uploaded_file.name.split(".")[-1].lower()
            if file_type == "txt":
                cv_text = uploaded_file.read().decode("utf-8")
                st.success("✅ Text file loaded!")
            elif file_type == "pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(uploaded_file)
                    cv_text = ""
                    for page in reader.pages:
                        cv_text += page.extract_text()
                    st.success(f"✅ PDF loaded — {len(reader.pages)} page(s) extracted!")
                except ImportError:
                    st.error("⚠️ pypdf not installed. Run: pip install pypdf")
            elif file_type == "docx":
                try:
                    import docx
                    doc = docx.Document(BytesIO(uploaded_file.read()))
                    cv_text = "\n".join([para.text for para in doc.paragraphs])
                    st.success("✅ Word document loaded!")
                except ImportError:
                    st.error("⚠️ python-docx not installed.")
            if cv_text:
                with st.expander("👀 Preview extracted CV text"):
                    st.text(cv_text[:500] + "..." if len(cv_text) > 500 else cv_text)
    else:
        cv_text = st.text_area(
            "Paste your CV text here",
            height=200,
            placeholder="Paste your full CV content here..."
        )

    st.markdown("---")

    st.markdown("**💼 Job Details**")
    col1, col2 = st.columns(2)
    with col1:
        role_title = st.text_input("Role Title *", placeholder="AI Engineering Intern")
        company_name = st.text_input("Company Name *", placeholder="Mercedes-Benz")
    with col2:
        company_address = st.text_input("Company Address (optional)", placeholder="Berlin, Germany")
        hiring_email = st.text_input("Hiring Manager Email (optional)", placeholder="hr@company.com")

    hiring_manager = st.text_input(
        "Hiring Manager Name (optional)",
        placeholder="Leave blank if unknown — will use 'Dear Hiring Team'"
    )
    job_description = st.text_area(
        "Paste Job Description *",
        height=200,
        placeholder="Paste the full job description here..."
    )

    st.markdown("---")

    if st.button("🚀 Generate My Application with ProofHire"):
        if not user_name or not user_email or not github_username \
                or not cv_text or not job_description \
                or not role_title or not company_name:
            st.error("⚠️ Please fill in all required fields marked with *")
        else:
            progress = st.progress(0)
            status = st.empty()
            try:
                initial_state = AgentState(
                    github_username=github_username,
                    cv_text=cv_text,
                    job_description=job_description,
                    github_analysis="",
                    cv_summary="",
                    match_analysis="",
                    cover_letter="",
                    user_name=user_name,
                    user_email=user_email,
                    user_phone=user_phone,
                    user_location=user_location,
                    user_linkedin=user_linkedin,
                    hiring_manager=hiring_manager,
                    company_name=company_name,
                    company_address=company_address,
                    hiring_email=hiring_email,
                    role_title=role_title
                )

                status.info("🔍 Agent 1: Fetching GitHub Profile...")
                progress.progress(10)
                app = build_graph()

                status.info("🔍 Agent 1: Analysing GitHub repositories...")
                progress.progress(25)
                result = app.invoke(initial_state)
                progress.progress(50)

                status.info("📄 Agent 2: Analysing CV...")
                progress.progress(60)

                status.info("🎯 Agent 3: Calculating Fit Score...")
                progress.progress(75)

                status.info("✍️ Agent 4: Writing Cover Letter...")
                progress.progress(90)

                status.info("💾 Saving to ProofHire database...")
                supabase.table("job_applications").insert({
                    "user_name": user_name,
                    "user_email": user_email,
                    "user_phone": user_phone,
                    "user_location": user_location,
                    "user_linkedin": user_linkedin,
                    "github_username": github_username,
                    "company_name": company_name,
                    "role_title": role_title,
                    "github_analysis": result['github_analysis'],
                    "cv_summary": result['cv_summary'],
                    "match_analysis": result['match_analysis'],
                    "cover_letter": result['cover_letter']
                }).execute()

                progress.progress(100)
                status.success("✅ Application Generated Successfully!")

                st.markdown("---")
                st.markdown(f"""
                    <div class="result-header">
                        <div class="result-title">🎯 Your ProofHire Results</div>
                        <div class="result-sub">{user_name} → {company_name} · {role_title}</div>
                    </div>
                """, unsafe_allow_html=True)

                with st.expander("🔍 GitHub Analysis", expanded=False):
                    st.markdown(result['github_analysis'])
                with st.expander("📄 CV Summary", expanded=False):
                    st.markdown(result['cv_summary'])
                with st.expander("🎯 Job Match Analysis & Technical Fit Score", expanded=True):
                    st.markdown(result['match_analysis'])

                st.markdown("### ✉️ Your Generated Cover Letter")
                st.text_area("Cover Letter", value=result['cover_letter'], height=400)

                word_buffer = create_word_doc(result['cover_letter'], user_name, company_name)
                full_results = (
                    f"GITHUB ANALYSIS:\n{result['github_analysis']}\n\n"
                    f"CV SUMMARY:\n{result['cv_summary']}\n\n"
                    f"JOB MATCH:\n{result['match_analysis']}\n\n"
                    f"COVER LETTER:\n{result['cover_letter']}"
                )

                st.markdown("**📥 Download Your Application**")
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button(
                        label="📄 Cover Letter (.docx)",
                        data=word_buffer,
                        file_name=f"ProofHire_CoverLetter_{company_name}_{user_name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with col2:
                    st.download_button(
                        label="📝 Cover Letter (.txt)",
                        data=result['cover_letter'],
                        file_name=f"ProofHire_CoverLetter_{company_name}_{user_name}.txt",
                        mime="text/plain"
                    )
                with col3:
                    st.download_button(
                        label="📊 Full Results (.txt)",
                        data=full_results,
                        file_name=f"ProofHire_FullResults_{company_name}_{user_name}.txt",
                        mime="text/plain"
                    )

            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
                progress.progress(0)

# ============================
# PAGE 3 — HISTORY
# ============================
elif page == "📚 My History":

    st.markdown("""
        <div style="margin-bottom:24px;">
            <div class="section-header">📚 Application History</div>
            <div class="section-sub">All your ProofHire applications saved in one place</div>
        </div>
    """, unsafe_allow_html=True)

    try:
        response = supabase.table("job_applications").select("*")\
            .order("created_at", desc=True).execute()
        records = response.data

        if not records:
            st.markdown("""
                <div style="background:white; border:1px solid #e2e8f0; border-radius:12px;
                            padding:40px; text-align:center;">
                    <div style="font-size:2rem; margin-bottom:12px;">📭</div>
                    <div style="font-weight:700; color:#0f172a; margin-bottom:8px;">No applications yet</div>
                    <div style="color:#64748b;">Click Generate Application above to create your first one!</div>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
                <div style="background:#eff6ff; border:1px solid #bfdbfe; border-radius:8px;
                            padding:12px 20px; margin-bottom:20px; color:#1d4ed8; font-weight:600;">
                    ✅ {len(records)} application(s) saved in ProofHire
                </div>
            """, unsafe_allow_html=True)

            for i, record in enumerate(records):
                with st.expander(
                    f"🏢 {record['company_name']} — {record['role_title']} "
                    f"| {record['user_name']} | {record['created_at'][:10]}"
                ):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**👤 Name:** {record['user_name']}")
                        st.markdown(f"**📧 Email:** {record['user_email']}")
                        st.markdown(f"**🐙 GitHub:** {record['github_username']}")
                    with col2:
                        st.markdown(f"**🏢 Company:** {record['company_name']}")
                        st.markdown(f"**💼 Role:** {record['role_title']}")
                        st.markdown(f"**📅 Date:** {record['created_at'][:10]}")

                    st.markdown("**✉️ Cover Letter:**")
                    st.text_area(
                        "Cover Letter",
                        value=record['cover_letter'],
                        height=300,
                        key=f"cl_{i}"
                    )

                    word_buffer_h = create_word_doc(
                        record['cover_letter'],
                        record['user_name'],
                        record['company_name']
                    )

                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="📄 Download Word (.docx)",
                            data=word_buffer_h,
                            file_name=f"ProofHire_{record['company_name']}_{record['user_name']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_word_{i}"
                        )
                    with col2:
                        st.download_button(
                            label="📝 Download Text (.txt)",
                            data=record['cover_letter'],
                            file_name=f"ProofHire_{record['company_name']}_{record['user_name']}.txt",
                            mime="text/plain",
                            key=f"dl_txt_{i}"
                        )

    except Exception as e:
        st.error(f"❌ Could not load history: {str(e)}")

# ---- FOOTER ----
st.markdown("""
    <div class="footer">
        🎯 ProofHire · GitHub-Verified Job Matching ·
        Powered by LangGraph + Groq + Supabase · 2026
    </div>
""", unsafe_allow_html=True)